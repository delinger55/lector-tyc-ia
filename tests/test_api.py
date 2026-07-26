"""Tests para app/routers/analyze.py y app/main.py.

Incluye Properties 1, 2, 9 y unit tests del endpoint.
"""

import io

import pytest
from docx import Document
from hypothesis import given, settings, HealthCheck
from hypothesis import strategies as st
from httpx import ASGITransport, AsyncClient
from pypdf import PdfWriter
from pypdf.generic import (
    DecodedStreamObject,
    DictionaryObject,
    NameObject,
)

from app.main import app


# --- Helpers para generar archivos de prueba ---


def _create_valid_pdf_bytes() -> bytes:
    """Genera un PDF válido con texto > 100 caracteres."""
    text = (
        "Este es un documento de terminos y condiciones de prueba "
        "con suficiente texto para superar el umbral de cien caracteres "
        "que necesitamos para validar correctamente el extractor."
    )
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    page = writer.pages[0]

    content = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET"
    stream = DecodedStreamObject()
    stream.set_data(content.encode("latin-1"))

    if "/Resources" not in page:
        page[NameObject("/Resources")] = DictionaryObject()
    resources = page["/Resources"]
    if "/Font" not in resources:
        resources[NameObject("/Font")] = DictionaryObject()
    font_dict = DictionaryObject()
    font_dict[NameObject("/Type")] = NameObject("/Font")
    font_dict[NameObject("/Subtype")] = NameObject("/Type1")
    font_dict[NameObject("/BaseFont")] = NameObject("/Helvetica")
    resources["/Font"][NameObject("/F1")] = font_dict
    page[NameObject("/Contents")] = writer._add_object(stream)

    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def _create_valid_docx_bytes() -> bytes:
    """Genera un DOCX válido con texto suficiente."""
    doc = Document()
    doc.add_paragraph(
        "Este es un documento de términos y condiciones de prueba. "
        "Contiene suficiente texto para ser procesado correctamente "
        "por el sistema de análisis con inteligencia artificial."
    )
    doc.add_paragraph(
        "El usuario acepta los términos descritos en este documento "
        "y reconoce que ha leído todas las cláusulas aquí contenidas."
    )
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _create_scanned_pdf_bytes() -> bytes:
    """Genera un PDF sin texto extraíble (simula escaneado)."""
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


# --- Fixture del cliente ---


@pytest.fixture
def anyio_backend():
    return "asyncio"


@pytest.fixture
async def client():
    """Cliente HTTP para tests de la API."""
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        yield ac


# --- Property 1: Validación de extensión ---


class TestProperty1ExtensionValidation:
    """Property 1: Solo .pdf y .docx son aceptados, resto → 400."""

    @settings(max_examples=50, deadline=None, suppress_health_check=[HealthCheck.function_scoped_fixture])
    @given(
        ext=st.sampled_from(
            [".doc", ".jpg", ".png", ".html", ".csv", ".xlsx", ".zip", ".exe", ".py", ".rtf"]
        )
    )
    async def test_invalid_extensions_return_400(self, ext: str, client: AsyncClient):
        """Extensiones no soportadas retornan 400 INVALID_FORMAT."""
        content = b"fake content for testing"
        filename = f"document{ext}"
        response = await client.post(
            "/api/v1/analyze",
            files={"file": (filename, content, "application/octet-stream")},
        )
        assert response.status_code == 400
        data = response.json()
        assert data["error_code"] == "INVALID_FORMAT"

    async def test_pdf_extension_accepted(self, client: AsyncClient):
        """Extensión .pdf es aceptada (no retorna INVALID_FORMAT)."""
        pdf_bytes = _create_valid_pdf_bytes()
        response = await client.post(
            "/api/v1/analyze",
            files={"file": ("test.pdf", pdf_bytes, "application/pdf")},
        )
        # No debe ser INVALID_FORMAT (puede ser otro error o 200)
        if response.status_code == 400:
            assert response.json()["error_code"] != "INVALID_FORMAT"

    async def test_docx_extension_accepted(self, client: AsyncClient):
        """Extensión .docx es aceptada (no retorna INVALID_FORMAT)."""
        docx_bytes = _create_valid_docx_bytes()
        response = await client.post(
            "/api/v1/analyze",
            files={"file": ("test.docx", docx_bytes, "application/octet-stream")},
        )
        if response.status_code == 400:
            assert response.json()["error_code"] != "INVALID_FORMAT"

    async def test_extension_case_insensitive(self, client: AsyncClient):
        """Extensión .PDF (mayúscula) es aceptada."""
        pdf_bytes = _create_valid_pdf_bytes()
        response = await client.post(
            "/api/v1/analyze",
            files={"file": ("test.PDF", pdf_bytes, "application/pdf")},
        )
        if response.status_code == 400:
            assert response.json()["error_code"] != "INVALID_FORMAT"


# --- Property 2: Validación de tamaño ---


class TestProperty2SizeValidation:
    """Property 2: Archivos que exceden MAX_FILE_SIZE → 400 FILE_TOO_LARGE."""

    async def test_file_exceeding_max_size_returns_400(self, client: AsyncClient):
        """Archivo > 10 MB retorna 400 FILE_TOO_LARGE."""
        # Crear contenido de ~11 MB (con extensión válida)
        large_content = b"X" * (11 * 1024 * 1024)
        response = await client.post(
            "/api/v1/analyze",
            files={"file": ("large.pdf", large_content, "application/pdf")},
        )
        assert response.status_code == 400
        data = response.json()
        assert data["error_code"] == "FILE_TOO_LARGE"

    async def test_file_within_size_not_rejected_for_size(self, client: AsyncClient):
        """Archivo < 10 MB no es rechazado por tamaño."""
        pdf_bytes = _create_valid_pdf_bytes()
        assert len(pdf_bytes) < 10 * 1024 * 1024  # Confirmar que es pequeño
        response = await client.post(
            "/api/v1/analyze",
            files={"file": ("test.pdf", pdf_bytes, "application/pdf")},
        )
        # No debe ser FILE_TOO_LARGE
        if response.status_code == 400:
            assert response.json()["error_code"] != "FILE_TOO_LARGE"


# --- Property 9: Archivos inválidos retornan HTTP 400 ---


class TestProperty9InvalidFilesReturn400:
    """Property 9: Archivos inválidos → HTTP 400 con error_code y detail en español."""

    async def test_unsupported_format_returns_400(self, client: AsyncClient):
        """Formato no soportado → 400."""
        response = await client.post(
            "/api/v1/analyze",
            files={"file": ("doc.rtf", b"contenido", "text/plain")},
        )
        assert response.status_code == 400
        data = response.json()
        assert "error_code" in data
        assert "detail" in data
        assert data["error_code"] == "INVALID_FORMAT"

    async def test_corrupt_pdf_returns_400(self, client: AsyncClient):
        """PDF corrupto → 400 CORRUPT_FILE."""
        corrupt = b"not a valid pdf at all" * 10
        response = await client.post(
            "/api/v1/analyze",
            files={"file": ("bad.pdf", corrupt, "application/pdf")},
        )
        assert response.status_code == 400
        data = response.json()
        assert data["error_code"] == "CORRUPT_FILE"

    async def test_scanned_pdf_returns_400(self, client: AsyncClient):
        """PDF escaneado (sin texto) → 400 SCANNED_DOCUMENT."""
        scanned = _create_scanned_pdf_bytes()
        response = await client.post(
            "/api/v1/analyze",
            files={"file": ("scan.pdf", scanned, "application/pdf")},
        )
        assert response.status_code == 400
        data = response.json()
        assert data["error_code"] == "SCANNED_DOCUMENT"

    async def test_error_response_has_spanish_detail(self, client: AsyncClient):
        """El campo detail está en español."""
        response = await client.post(
            "/api/v1/analyze",
            files={"file": ("doc.rtf", b"x", "text/plain")},
        )
        data = response.json()
        # Verificar que el mensaje contiene palabras en español
        assert "PDF" in data["detail"] or "Word" in data["detail"]


# --- Unit tests del endpoint ---


class TestAnalyzeEndpointUnit:
    """Unit tests para el endpoint POST /api/v1/analyze."""

    async def test_get_home_returns_200_html(self, client: AsyncClient):
        """GET / retorna 200 con content-type HTML."""
        response = await client.get("/")
        assert response.status_code == 200
        assert "text/html" in response.headers["content-type"]

    async def test_get_home_contains_title(self, client: AsyncClient):
        """GET / contiene el título de la aplicación."""
        response = await client.get("/")
        assert "Lector de Términos y Condiciones" in response.text

    async def test_valid_pdf_returns_200_with_analysis(self, client: AsyncClient):
        """PDF válido → 200 con AnalysisResult JSON."""
        pdf_bytes = _create_valid_pdf_bytes()
        response = await client.post(
            "/api/v1/analyze",
            files={"file": ("terms.pdf", pdf_bytes, "application/pdf")},
        )
        assert response.status_code == 200
        data = response.json()
        assert "is_valid_terms" in data
        assert "summary_points" in data
        assert "risk_clauses" in data
        assert len(data["summary_points"]) == 5

    async def test_valid_docx_returns_200_with_analysis(self, client: AsyncClient):
        """DOCX válido → 200 con AnalysisResult JSON."""
        docx_bytes = _create_valid_docx_bytes()
        response = await client.post(
            "/api/v1/analyze",
            files={"file": ("terms.docx", docx_bytes, "application/octet-stream")},
        )
        assert response.status_code == 200
        data = response.json()
        assert "is_valid_terms" in data
        assert len(data["summary_points"]) == 5

    async def test_response_json_structure(self, client: AsyncClient):
        """Respuesta exitosa tiene la estructura completa de AnalysisResult."""
        pdf_bytes = _create_valid_pdf_bytes()
        response = await client.post(
            "/api/v1/analyze",
            files={"file": ("doc.pdf", pdf_bytes, "application/pdf")},
        )
        assert response.status_code == 200
        data = response.json()
        # Campos obligatorios
        assert isinstance(data["is_valid_terms"], bool)
        assert isinstance(data["summary_points"], list)
        assert isinstance(data["risk_clauses"], list)
        # Cada cláusula tiene estructura correcta
        if data["risk_clauses"]:
            clause = data["risk_clauses"][0]
            assert "title" in clause
            assert "severity" in clause
            assert "explanation" in clause
            assert clause["severity"] in ["HIGH", "MEDIUM", "LOW"]


# --- Tests para .txt y URL ---


class TestTxtFileEndpoint:
    """Tests para el flujo de archivos .txt."""

    async def test_valid_txt_returns_200(self, client: AsyncClient):
        """Archivo .txt válido con texto suficiente → 200."""
        content = ("Este es un documento de prueba con términos y condiciones. " * 5).encode("utf-8")
        response = await client.post(
            "/api/v1/analyze",
            files={"file": ("terms.txt", content, "text/plain")},
        )
        assert response.status_code == 200
        data = response.json()
        assert len(data["summary_points"]) == 5

    async def test_txt_extension_accepted(self, client: AsyncClient):
        """Extensión .txt no retorna INVALID_FORMAT."""
        content = ("Texto de prueba suficiente para el análisis del documento. " * 3).encode("utf-8")
        response = await client.post(
            "/api/v1/analyze",
            files={"file": ("readme.txt", content, "text/plain")},
        )
        if response.status_code == 400:
            assert response.json()["error_code"] != "INVALID_FORMAT"


class TestUrlEndpoint:
    """Tests para el flujo de análisis por URL."""

    async def test_url_without_file_is_accepted(self, client: AsyncClient):
        """Enviar solo URL (sin file) es procesado por el endpoint."""
        # Usamos una URL inválida para verificar que llega al extractor
        response = await client.post(
            "/api/v1/analyze",
            data={"url": "not-a-valid-url"},
        )
        # Debe retornar 400 URL_EXTRACTION_FAILED (URL inválida), no 422 (validation)
        assert response.status_code == 400
        data = response.json()
        assert data["error_code"] == "URL_EXTRACTION_FAILED"

    async def test_url_invalid_protocol_returns_400(self, client: AsyncClient):
        """URL sin http/https → 400."""
        response = await client.post(
            "/api/v1/analyze",
            data={"url": "ftp://ejemplo.com/terms"},
        )
        assert response.status_code == 400

    async def test_url_empty_returns_400(self, client: AsyncClient):
        """URL vacía con file vacío → 400."""
        response = await client.post(
            "/api/v1/analyze",
            data={"url": ""},
        )
        assert response.status_code == 400
