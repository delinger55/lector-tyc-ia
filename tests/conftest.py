"""Fixtures compartidas para la suite de tests.

Provee: AsyncClient, archivos de prueba generados programáticamente,
y configuración de hypothesis.
"""

import io

import pytest
from docx import Document
from httpx import ASGITransport, AsyncClient
from hypothesis import settings as hypothesis_settings
from pypdf import PdfWriter
from pypdf.generic import (
    DecodedStreamObject,
    DictionaryObject,
    NameObject,
)

from app.main import app

# --- Configuración global de hypothesis ---
hypothesis_settings.register_profile(
    "default", max_examples=200, deadline=None
)
hypothesis_settings.load_profile("default")


# --- Fixture: AsyncClient ---


@pytest.fixture
async def client():
    """Cliente HTTP async para tests de integración."""
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        yield ac


# --- Fixtures: Archivos de prueba ---


@pytest.fixture
def sample_pdf_bytes() -> bytes:
    """PDF válido con texto > 100 caracteres (extraíble por pypdf)."""
    text = (
        "Este es un documento de terminos y condiciones de prueba. "
        "Contiene suficiente texto para superar el umbral de cien caracteres "
        "que el sistema requiere para considerar el documento como valido y procesable."
    )
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    page = writer.pages[0]

    content = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET"
    stream = DecodedStreamObject()
    stream.set_data(content.encode("latin-1"))

    if "/Resources" not in page:
        page[NameObject("/Resources")] = DictionaryObject()
    resources = page["/Resources"]
    if "/Font" not in resources:
        resources[NameObject("/Font")] = DictionaryObject()
    font_dict = DictionaryObject()
    font_dict[NameObject("/Type")] = NameObject("/Font")
    font_dict[NameObject("/Subtype")] = NameObject("/Type1")
    font_dict[NameObject("/BaseFont")] = NameObject("/Helvetica")
    resources["/Font"][NameObject("/F1")] = font_dict
    page[NameObject("/Contents")] = writer._add_object(stream)

    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


@pytest.fixture
def sample_docx_bytes() -> bytes:
    """DOCX válido con múltiples párrafos de texto."""
    doc = Document()
    doc.add_paragraph(
        "Este es un documento de términos y condiciones de prueba. "
        "Contiene suficiente texto para ser procesado correctamente "
        "por el sistema de análisis con inteligencia artificial."
    )
    doc.add_paragraph(
        "El usuario acepta los términos descritos en este documento "
        "y reconoce que ha leído todas las cláusulas aquí contenidas."
    )
    doc.add_paragraph(
        "La empresa se reserva el derecho de modificar estos términos "
        "en cualquier momento sin notificación previa al usuario."
    )
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def scanned_pdf_bytes() -> bytes:
    """PDF sin texto extraíble (simula documento escaneado, < 100 chars)."""
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


@pytest.fixture
def large_pdf_bytes() -> bytes:
    """PDF que excede el tamaño máximo (> 10 MB)."""
    # Generar contenido de ~11 MB con extensión PDF válida
    return b"%PDF-1.4\n" + (b"X" * (11 * 1024 * 1024))


@pytest.fixture
def corrupt_pdf_bytes() -> bytes:
    """Bytes que no son un PDF válido pero tienen extensión .pdf."""
    return b"This is definitely not a valid PDF file content at all" * 20
